#!/usr/bin/env python3
"""
Translate a .docx file EN→DE via DeepL document translation API.
Formatting, styles, tables, and footnotes are preserved by DeepL natively.

Usage:
    uv run python translate_docx.py --dry-run          # quota check only
    uv run python translate_docx.py                    # translate
    uv run python translate_docx.py --input other.docx --output other_DE.docx
    uv run python translate_docx.py --formality more   # formal "Sie" (default: less)
"""

import argparse
import os
import sys
from pathlib import Path

import deepl
from docx import Document
from dotenv import load_dotenv

load_dotenv()

# ── Defaults ──────────────────────────────────────────────────────────────────

DEFAULT_INPUT  = Path("scratch/Abundance_from_epub_260609_TOC.docx")
DEFAULT_OUTPUT = Path("scratch/Abundance_from_epub_260609_TOC_DE.docx")
SOURCE_LANG    = "EN"
TARGET_LANG    = "DE"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _translator() -> deepl.Translator:
    key = os.environ.get("DEEPL_API_KEY", "")
    if not key:
        sys.exit("ERROR: DEEPL_API_KEY not set in .env")
    return deepl.Translator(key)


def _count_chars(path: Path) -> int:
    """Approximate billable character count by summing all text in the docx."""
    doc = Document(str(path))
    total = sum(len(p.text) for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += sum(len(p.text) for p in cell.paragraphs)
    return total


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="Translate .docx EN→DE via DeepL")
    parser.add_argument("--input",    default=str(DEFAULT_INPUT),  help="Source .docx")
    parser.add_argument("--output",   default=str(DEFAULT_OUTPUT), help="Output .docx")
    parser.add_argument("--formality", default="less", choices=["less", "more", "default"],
                        help="less = Du, more = Sie (default: less)")
    parser.add_argument("--dry-run",  action="store_true", help="Quota check only, no API call")
    args = parser.parse_args()

    input_path  = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        sys.exit(f"ERROR: file not found: {input_path}")

    t = _translator()

    # Quota + estimate
    usage     = t.get_usage()
    remaining = usage.character.limit - usage.character.count
    estimated = _count_chars(input_path)
    file_kb   = input_path.stat().st_size / 1024
    glossary_id = os.environ.get("DEEPL_GLOSSARY_ID", "")

    print(f"Input:           {input_path}  ({file_kb:.0f} KB)")
    print(f"Output:          {output_path}")
    print(f"Estimated chars: {estimated:,}")
    print(f"Quota used:      {usage.character.count:,} / {usage.character.limit:,}")
    print(f"Quota remaining: {remaining:,}")
    print(f"Glossary:        {glossary_id or '(none)'}")
    print(f"Formality:       {args.formality}")

    if estimated > remaining:
        print("\nWARNING: estimated chars exceed remaining quota — translation will likely fail.")

    if args.dry_run:
        print("\nDry run — no API call made.")
        return

    # Translate
    output_path.parent.mkdir(parents=True, exist_ok=True)
    kwargs: dict = dict(source_lang=SOURCE_LANG, target_lang=TARGET_LANG, formality=args.formality)
    if glossary_id:
        kwargs["glossary"] = glossary_id

    print("\nTranslating — this may take a minute for large documents…")
    t.translate_document_from_filepath(str(input_path), str(output_path), **kwargs)

    usage2 = t.get_usage()
    used = usage2.character.count - usage.character.count
    print(f"Done.  Chars used: {used:,}  |  Remaining: {usage2.character.limit - usage2.character.count:,}")
    print(f"Output: {output_path.resolve()}")


if __name__ == "__main__":
    main()
